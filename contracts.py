import customtkinter as ctk
from tkinter import StringVar, BooleanVar, messagebox, filedialog
from pathlib import Path
from datetime import datetime, timedelta

from docx import Document
import requests
from num2words import num2words
import re
import sys

# Diretórios base (funciona tanto no script quanto empacotado com PyInstaller)
if getattr(sys, "frozen", False):
    # Rodando empacotado (PyInstaller): arquivos estão em sys._MEIPASS
    BASE_DIR = Path(sys._MEIPASS)
else:
    # Rodando via `python contracts.py`
    BASE_DIR = Path(__file__).parent

APP_NAME = "Contratos Musicais"
APP_VERSION = "0.1.0"  # sem o 'v'

# caminhos dos templates (ajusta se a pasta for outra)
TEMPLATES_DIR = BASE_DIR / "templates"
TEMPLATE_CONTRATO = TEMPLATES_DIR / "contrato_som_banda.docx"

# pasta de saída para contratos gerados
SAIDA_DIR = Path.cwd() / "contratos_gerados"
SAIDA_DIR.mkdir(exist_ok=True)

def valor_por_extenso(valor: str) -> str:
    """
    Converte uma string de valor monetário brasileiro (ex.: '2000', '2.000,00', 'R$ 2.000,00')
    em texto por extenso, ex.: 'dois mil reais' ou 'dois mil reais e cinquenta centavos'.
    """
    if not valor:
        return ""

    v = (
        valor.replace("R$", "")
        .replace(" ", "")
        .replace(".", "")
        .replace(",", ".")
    )
    try:
        numero = float(v)
    except Exception:
        return ""

    reais = int(numero)
    centavos = int(round((numero - reais) * 100))

    extenso_reais = num2words(reais, lang="pt_BR")
    # ajustes cosméticos simples
    extenso_reais = extenso_reais.replace(" e zero", "")

    if centavos == 0:
        return f"{extenso_reais} reais"
    else:
        extenso_cent = num2words(centavos, lang="pt_BR")
        return f"{extenso_reais} reais e {extenso_cent} centavos"


# ------------------------------------------------------------------
# Utilitários de horários por extenso
# ------------------------------------------------------------------
def hora_por_extenso(hora: int, minuto: int) -> str:
    """
    Converte hora e minuto em texto por extenso, ex.:
    18:00 -> 'dezoito horas'
    18:30 -> 'dezoito horas e trinta minutos'
    """
    if hora < 0 or hora > 23 or minuto < 0 or minuto > 59:
        return ""

    # Ajuste de gênero para 'hora' (feminino)
    if hora == 1:
        hora_ext = "uma"
    elif hora == 2:
        hora_ext = "duas"
    else:
        hora_ext = num2words(hora, lang="pt_BR")

    if hora == 1:
        base = f"{hora_ext} hora"
    else:
        base = f"{hora_ext} horas"

    if minuto == 0:
        return base

    min_ext = num2words(minuto, lang="pt_BR")
    if minuto == 1:
        return f"{base} e {min_ext} minuto"
    else:
        return f"{base} e {min_ext} minutos"


def parse_hora_minuto(texto: str):
    """
    Tenta interpretar uma string hh:mm e retorna (hora, minuto) ou None se inválido.
    """
    if not texto:
        return None
    m = re.match(r"^\s*(\d{1,2}):(\d{2})\s*$", texto)
    if not m:
        return None
    hora = int(m.group(1))
    minuto = int(m.group(2))
    if 0 <= hora <= 23 and 0 <= minuto <= 59:
        return hora, minuto
    return None

def data_por_extenso(data_str: str) -> str:
    """
    Converte 'dd/mm/aaaa' para 'dd de Mês de aaaa', ex.: '06/01/2025' -> '06 de Janeiro de 2025'.
    Se não conseguir interpretar, retorna a string original.
    """
    if not data_str:
        return ""
    m = re.match(r"^\s*(\d{1,2})/(\d{1,2})/(\d{4})\s*$", data_str)
    if not m:
        return data_str

    dia = int(m.group(1))
    mes = int(m.group(2))
    ano = int(m.group(3))

    meses = {
        1: "Janeiro",
        2: "Fevereiro",
        3: "Março",
        4: "Abril",
        5: "Maio",
        6: "Junho",
        7: "Julho",
        8: "Agosto",
        9: "Setembro",
        10: "Outubro",
        11: "Novembro",
        12: "Dezembro",
    }

    nome_mes = meses.get(mes)
    if not nome_mes:
        return data_str

    return f"{dia:02d} de {nome_mes} de {ano}"

def montar_contexto(values: dict, som: str, alimentacao: str) -> dict:
    """Monta o dicionário de placeholders -> valores para usar no DOCX."""

    # ------------------------------------------------------------------
    # CONTRATANTE
    # ------------------------------------------------------------------
    endereco_contratante = (
        f"{values.get('contratante_endereco_logradouro', '')}, "
        f"{values.get('contratante_endereco_numero', '')} "
        f"{values.get('contratante_endereco_complemento', '')} - "
        f"{values.get('contratante_endereco_bairro', '')}, "
        f"{values.get('contratante_endereco_cidade', '')}/"
        f"{values.get('contratante_endereco_uf', '')} - "
        f"CEP {values.get('contratante_endereco_cep', '')}"
    )

    # ------------------------------------------------------------------
    # CONTRATADO
    # ------------------------------------------------------------------
    endereco_contratado = (
        f"{values.get('contratado_endereco_logradouro', '')}, "
        f"{values.get('contratado_endereco_numero', '')} "
        f"{values.get('contratado_endereco_complemento', '')} - "
        f"{values.get('contratado_endereco_bairro', '')}, "
        f"{values.get('contratado_endereco_cidade', '')}/"
        f"{values.get('contratado_endereco_uf', '')} - "
        f"CEP {values.get('contratado_endereco_cep', '')}"
    )

    # ------------------------------------------------------------------
    # EVENTO
    # ------------------------------------------------------------------
    horario = (
        f"{values.get('evento_horario_inicio', '')}h às "
        f"{values.get('evento_horario_fim_previsto', '')}h"
    )
    
    evento_data_raw = values.get("evento_data", "")
    evento_data_ext = data_por_extenso(evento_data_raw)

    nome_local = values.get("evento_local_nome", "")
    logradouro_local = values.get("evento_local_logradouro", "")
    numero_local = values.get("evento_local_numero", "").strip()
    compl_local = values.get("evento_local_complemento", "").strip()
    bairro_local = values.get("evento_local_bairro", "")
    cidade_local = values.get("evento_local_cidade", "")
    uf_local = values.get("evento_local_uf", "")
    cep_local = values.get("evento_local_cep", "")

    partes = [nome_local, logradouro_local]

    if numero_local:
        partes.append(numero_local)
    if compl_local:
        # complemento logo após o número, sem vírgula
        if partes:
            partes[-1] = f"{partes[-1]} {compl_local}"
        else:
            partes.append(compl_local)

    endereco_evento = ", ".join(p for p in partes if p)
    endereco_evento += f" - {bairro_local}, {cidade_local}/{uf_local} - CEP {cep_local}"

    # ------------------------------------------------------------------
    # DURAÇÃO DO EVENTO E HORÁRIO DE CHEGADA
    # ------------------------------------------------------------------
    evento_duracao_str = ""
    evento_horario_chegada_str = ""

    inicio_str = values.get("evento_horario_inicio", "").strip()
    fim_str = values.get("evento_horario_fim_previsto", "").strip()

    inicio_parsed = parse_hora_minuto(inicio_str)
    fim_parsed = parse_hora_minuto(fim_str)

    if inicio_parsed and fim_parsed:
        ih, im = inicio_parsed
        fh, fm = fim_parsed

        # usa uma data fictícia apenas para cálculo de diferença
        inicio_dt = datetime(2000, 1, 1, ih, im)
        fim_dt = datetime(2000, 1, 1, fh, fm)
        if fim_dt <= inicio_dt:
            # se o fim for menor ou igual ao início, assume virada de dia
            fim_dt += timedelta(days=1)

        diff = fim_dt - inicio_dt
        total_min = diff.seconds // 60
        horas = total_min // 60
        minutos = total_min % 60

        # duração formatada: HH:MM (por extenso)
        dur_num = f"{horas:02d}:{minutos:02d}"
        dur_ext = hora_por_extenso(horas, minutos)
        evento_duracao_str = f"{dur_num} ({dur_ext})" if dur_ext else dur_num

        # horário de chegada: 1h antes do início
        chegada_dt = inicio_dt - timedelta(hours=1)
        ch_h = chegada_dt.hour
        ch_m = chegada_dt.minute
        ch_num = chegada_dt.strftime("%H:%M")
        ch_ext = hora_por_extenso(ch_h, ch_m)
        evento_horario_chegada_str = f"{ch_num} ({ch_ext})" if ch_ext else ch_num

    # ------------------------------------------------------------------
    # PAGAMENTO
    # ------------------------------------------------------------------
    forma = values.get("pagamento_forma", "")
    meio = values.get("pagamento_meio", "")
    valor_total = values.get("pagamento_valor_total", "").strip()

    valor_extenso = valor_por_extenso(valor_total)

    forma_pag = forma
    meio_pag = meio
    valor_num = valor_total

    pagamento_texto = ""

    if forma_pag == "À vista":
        data_unica = values.get("pagamento_data_unica", "")
        data_unica_ext = data_por_extenso(data_unica)
        pagamento_texto = (
            f"O pagamento será efetuado à vista, no valor total de R$ {valor_num} "
            f"({valor_extenso}), na data de {data_unica_ext}, via {meio_pag}."
        )

    elif forma_pag == "Sinal + restante":
        sinal = values.get("pagamento_sinal_percentual", "")
        data_sinal = values.get("pagamento_sinal_data", "")
        data_restante = values.get("pagamento_restante_data", "")
        data_sinal_ext = data_por_extenso(data_sinal)
        data_restante_ext = data_por_extenso(data_restante)

        sinal_info = f"{sinal}%"
        try:
            v_clean = valor_total.replace("R$", "").replace(" ", "").replace(".", "").replace(",", ".")
            total_float = float(v_clean)
            sinal_percent = float(str(sinal).replace(",", ".") or "0")
            valor_sinal = total_float * sinal_percent / 100.0

            tmp = f"{valor_sinal:,.2f}"  # 1,234.56
            tmp = tmp.replace(",", "X").replace(".", ",").replace("X", ".")
            valor_sinal_num = f"R$ {tmp}"
            valor_sinal_extenso = valor_por_extenso(tmp)

            sinal_info = f"{sinal}%, equivalente a {valor_sinal_num} ({valor_sinal_extenso}),"
        except Exception:
            pass

        pagamento_texto = (
            f"O pagamento será realizado em duas etapas: sinal de {sinal_info} até a data {data_sinal_ext} "
            f"e o valor restante até a data {data_restante_ext}, totalizando R$ {valor_num} "
            f"({valor_extenso}), via {meio_pag}."
        )

    elif forma_pag == "Parcelado":
        parcelas = values.get("pagamento_num_parcelas", "")
        primeira = values.get("pagamento_primeira_parcela_data", "")
        periodicidade = values.get("pagamento_periodicidade", "")
        primeira_ext = data_por_extenso(primeira)
        pagamento_texto = (
            f"O pagamento será efetuado em {parcelas} parcelas {periodicidade.lower()}, "
            f"a primeira com vencimento em {primeira_ext}, totalizando R$ {valor_num} "
            f"({valor_extenso}), via {meio_pag}."
        )

    else:  # Outro
        pagamento_texto = (
            f"O pagamento será realizado no valor total de R$ {valor_num} "
            f"({valor_extenso}), conforme forma negociada entre as partes, via {meio_pag}."
        )

    # Descrição simples (mantida para retrocompatibilidade se quiser usar)
    pagamento_descr = (
        f"O valor total de R$ {valor_num} será pago na forma '{forma_pag}', por meio de {meio_pag}."
    )

    # ------------------------------------------------------------------
    # FAVORECIDO
    # ------------------------------------------------------------------
    pix = values.get("favorecido_pix_chave", "")
    pix_tipo = values.get("favorecido_pix_tipo", "")
    pix_descr = f"{pix} ({pix_tipo})" if pix or pix_tipo else ""

    banco = values.get("favorecido_banco_nome", "")
    codigo = values.get("favorecido_banco_codigo", "")
    agencia = values.get("favorecido_agencia", "")
    conta = values.get("favorecido_conta", "")
    tipo_conta = values.get("favorecido_tipo_conta", "")

    partes_banco = []

    if banco:
        if codigo:
            partes_banco.append(f"Banco {banco} (Código {codigo})")
        else:
            partes_banco.append(f"Banco {banco}")

    if agencia:
        partes_banco.append(f"Agência {agencia}")

    if conta:
        partes_banco.append(f"Conta {conta}")

    if tipo_conta:
        partes_banco.append(f"– Conta {tipo_conta}")

    favorecido_dados_bancarios = ", ".join(partes_banco) if partes_banco else ""

    # ------------------------------------------------------------------
    # SOM - cláusula dinâmica
    # ------------------------------------------------------------------
    if som == "Banda":
        som_clausula = (
            "A banda CONTRATADA será responsável por levar, montar e operar o sistema de som "
            "necessário para a execução do show, incluindo mesa de som, amplificação, microfones "
            "e demais equipamentos de áudio, em condições adequadas ao ambiente do evento."
        )
    else:
        som_clausula = (
            "O CONTRATANTE será responsável por fornecer, montar e operar o sistema de som "
            "necessário para a execução do show, incluindo mesa de som, amplificação, microfones "
            "e demais equipamentos de áudio, em condições adequadas ao ambiente do evento."
        )
    
    # ------------------------------------------------------------------
    # ALIMENTAÇÃO - cláusula opcional
    # ------------------------------------------------------------------
    if alimentacao == "Sim":
        alimentacao_clausula = (
            "Cláusula 5.6. Fornecer consumação de alimentos e bebidas ao staff da banda no buffet "
            "presente do evento, caso haja buffet contratado. Em caso de comercialização de alimentos "
            "e bebidas no local do evento, as despesas decorrentes da consumação do(a) CONTRATADO(A) "
            "durante a apresentação artística correrão por conta do(a) CONTRATANTE até o limite de "
            "R$ 500,00 (quinhentos reais), caso esse seja ultrapassado as despesas serão de "
            "responsabilidade do(a) CONTRATADO(A)."
        )
    else:
        alimentacao_clausula = ""
    
    # ------------------------------------------------------------------
    # DATA DO CONTRATO (data corrente por extenso)
    # ------------------------------------------------------------------
    hoje_str = datetime.now().strftime("%d/%m/%Y")
    data_contrato_str = data_por_extenso(hoje_str)
    
    # ------------------------------------------------------------------
    # CONTEXTO BASE
    # ------------------------------------------------------------------
    contexto = {
        # CONTRATANTE
        "CONTRATANTE_NOME": values.get("contratante_nome_razao", ""),
        "CONTRATANTE_CPF_CNPJ": values.get("contratante_cpf_cnpj", ""),
        "CONTRATANTE_ENDERECO_COMPLETO": endereco_contratante,
        "CONTRATANTE_TELEFONE": values.get("contratante_telefone", ""),
        "CONTRATANTE_EMAIL": values.get("contratante_email", ""),

        # CONTRATADO
        "CONTRATADO_TIPO": values.get("contratado_tipo", ""),
        "CONTRATADO_NOME": values.get("contratado_nome_razao", ""),
        "CONTRATADO_CPF_CNPJ": values.get("contratado_cpf_cnpj", ""),
        "CONTRATADO_TELEFONE": values.get("contratado_telefone", ""),
        "CONTRATADO_EMAIL": values.get("contratado_email", ""),
        "CONTRATADO_ENDERECO_COMPLETO": endereco_contratado,
        "CONTRATADO_REPRESENTANTE_NOME": values.get("contratado_representante_nome", ""),
        "CONTRATADO_REPRESENTANTE_CPF": values.get("contratado_representante_cpf", ""),

        # EVENTO
        "EVENTO_NOME": values.get("evento_nome", ""),
        "ATRACAO_MUSICAL": values.get("evento_atracao_musical", ""),
        "EVENTO_DATA": evento_data_ext,
        "EVENTO_HORARIO": horario,
        "EVENTO_LOCAL_COMPLETO": endereco_evento,
        "EVENTO_DURACAO": evento_duracao_str,
        "EVENTO_HORARIO_CHEGADA": evento_horario_chegada_str,

        # PAGAMENTO
        "PAGAMENTO_VALOR_TOTAL": valor_num,
        "PAGAMENTO_VALOR_TOTAL_EXTENSO": valor_extenso,
        "PAGAMENTO_DESCRICAO": pagamento_descr,
        "PAGAMENTO_FORMA_DESCRICAO": pagamento_texto,

        # FAVORECIDO
        "FAVORECIDO_NOME": values.get("favorecido_nome", ""),
        "FAVORECIDO_CPF_CNPJ": values.get("favorecido_cpf_cnpj", ""),
        "FAVORECIDO_PIX": pix_descr,
        "FAVORECIDO_DADOS_BANCARIOS": favorecido_dados_bancarios,
        
        # SOM / ALIMENTAÇÃO
        "SOM_CLAUSULA": som_clausula,
        "ALIMENTACAO": alimentacao_clausula,
        
        # DATA DO CONTRATO
        "DATA_CONTRATO": data_contrato_str,
    }

    return contexto


def preencher_template_docx(caminho_template: Path, caminho_saida: Path, contexto: dict):
    """Abre o template DOCX, troca placeholders {{CHAVE}} pelos valores e salva no caminho de saída."""
    doc = Document(str(caminho_template))

    # Substitui em parágrafos
    for p in doc.paragraphs:
        for chave, valor in contexto.items():
            token = "{{" + chave + "}}"
            if token in p.text:
                p.text = p.text.replace(token, str(valor))

    # Substitui também em tabelas (muitos contratos usam tabelas)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for chave, valor in contexto.items():
                    token = "{{" + chave + "}}"
                    if token in cell.text:
                        cell.text = cell.text.replace(token, str(valor))

    doc.save(str(caminho_saida))


class ContractApp(ctk.CTk):
    def __init__(self):
        super().__init__()

        # garantir que a janela vem pra frente
        self.lift()
        self.attributes("-topmost", True)
        self.after(200, lambda: self.attributes("-topmost", False))

        ctk.set_appearance_mode("system")
        ctk.set_default_color_theme("blue")

        self.title(f"{APP_NAME} – v{APP_VERSION}")
        self.geometry("1100x700")

        # dicionário para guardar referências dos campos
        self.inputs = {}
        self.som_responsavel_var = StringVar(value="Contratante")
        self.alimentacao_var = StringVar(value="Sim")
        self.favorecido_igual_contratado_var = BooleanVar(value=False)
        self.pag_frame_avista = None
        self.pag_frame_sinal = None
        self.pag_frame_parc = None

        # --- TABVIEW PRINCIPAL ---
        self.tabview = ctk.CTkTabview(self)
        self.tabview.pack(fill="both", expand=True, padx=20, pady=(20, 10))
        self.tabview.configure(command=self._on_tab_change)

        self._build_tabs()
        self._setup_masks()
        self._update_pagamento_forma_ui("À vista")

        # --- RODAPÉ COM BOTÕES ---
        btn_frame = ctk.CTkFrame(self)
        btn_frame.pack(fill="x", padx=20, pady=(0, 15))
        
        btn_carregar = ctk.CTkButton(btn_frame, text="Carregar preenchimento", command=self.carregar_preenchimento)
        btn_carregar.pack(side="left", padx=(0, 10))

        btn_gerar = ctk.CTkButton(btn_frame, text="Gerar contrato", command=self.gerar_contrato)
        btn_gerar.pack(side="left", padx=(0, 10))

        btn_limpar = ctk.CTkButton(btn_frame, text="Limpar", fg_color="gray", command=self.limpar_campos)
        btn_limpar.pack(side="left", padx=(0, 10))

        btn_sair = ctk.CTkButton(btn_frame, text="Sair", fg_color="red", command=self.destroy)
        btn_sair.pack(side="right")

    # ---------------------------------------------------------
    # Construção das abas
    # ---------------------------------------------------------
    def _build_tabs(self):
        tab_contratante = self.tabview.add("Contratante")
        tab_contratado = self.tabview.add("Contratado")
        tab_evento = self.tabview.add("Evento / Local")
        tab_som = self.tabview.add("Som")
        tab_pagamento = self.tabview.add("Pagamento")
        tab_favorecido = self.tabview.add("Favorecido")
        tab_resumo = self.tabview.add("Resumo")

        self._build_tab_contratante(tab_contratante)
        self._build_tab_contratado(tab_contratado)
        self._build_tab_evento(tab_evento)
        self._build_tab_som(tab_som)
        self._build_tab_pagamento(tab_pagamento)
        self._build_tab_favorecido(tab_favorecido)
        self._build_tab_resumo(tab_resumo)

    def _setup_masks(self):
        """Configura máscaras de entrada para alguns campos específicos."""
        # Telefones
        self._attach_mask("contratante_telefone", "phone")
        self._attach_mask("contratado_telefone", "phone")

        # CPF/CNPJ principais
        self._attach_mask("contratante_cpf_cnpj", "cpf_cnpj")
        self._attach_mask("contratado_cpf_cnpj", "cpf_cnpj")
        self._attach_mask("favorecido_cpf_cnpj", "cpf_cnpj")

        # CPF de representantes
        self._attach_mask("contratante_representante_cpf", "cpf_cnpj")
        self._attach_mask("contratado_representante_cpf", "cpf_cnpj")

        # CEPs
        self._attach_mask("contratante_endereco_cep", "cep")
        self._attach_mask("contratado_endereco_cep", "cep")
        self._attach_mask("evento_local_cep", "cep")

        # Datas (dd/mm/aaaa)
        self._attach_mask("evento_data", "date")
        self._attach_mask("pagamento_data_unica", "date")
        self._attach_mask("pagamento_sinal_data", "date")
        self._attach_mask("pagamento_restante_data", "date")
        self._attach_mask("pagamento_primeira_parcela_data", "date")
        
        # Horários (hh:mm)
        self._attach_mask("evento_horario_inicio", "time")
        self._attach_mask("evento_horario_fim_previsto", "time")
        
        # Valor monetário˜
        self._attach_mask("pagamento_valor_total", "money")

    def _attach_mask(self, key: str, kind: str):
        """Anexa uma máscara de digitação ao campo identificado por 'key'."""
        entry = self.inputs.get(key)
        if not entry:
            return

        def on_key_release(event):
            text = entry.get()
            digits = re.sub(r"\D", "", text)

            if kind == "phone":
                # telefone: (##) #####-####
                digits_trim = digits[:11]
                if len(digits_trim) <= 2:
                    formatted = f"({digits_trim}"
                elif len(digits_trim) <= 7:
                    formatted = f"({digits_trim[:2]}) {digits_trim[2:]}"
                elif len(digits_trim) <= 10:
                    # formato (##) ####-####
                    formatted = f"({digits_trim[:2]}) {digits_trim[2:6]}-{digits_trim[6:]}"
                else:
                    # 11 dígitos: (##) #####-####
                    formatted = f"({digits_trim[:2]}) {digits_trim[2:7]}-{digits_trim[7:]}"
            elif kind == "cep":
                # CEP: #####-###
                digits_trim = digits[:8]
                if len(digits_trim) <= 5:
                    formatted = digits_trim
                else:
                    formatted = f"{digits_trim[:5]}-{digits_trim[5:]}"
            elif kind == "cpf_cnpj":
                # decide entre CPF e CNPJ pela quantidade de dígitos
                if len(digits) <= 11:
                    d = digits[:11]
                    if len(d) <= 3:
                        formatted = d
                    elif len(d) <= 6:
                        formatted = f"{d[:3]}.{d[3:]}"
                    elif len(d) <= 9:
                        formatted = f"{d[:3]}.{d[3:6]}.{d[6:]}"
                    else:
                        formatted = f"{d[:3]}.{d[3:6]}.{d[6:9]}-{d[9:]}"
                else:
                    d = digits[:14]
                    if len(d) <= 2:
                        formatted = d
                    elif len(d) <= 5:
                        formatted = f"{d[:2]}.{d[2:]}"
                    elif len(d) <= 8:
                        formatted = f"{d[:2]}.{d[2:5]}.{d[5:]}"
                    elif len(d) <= 12:
                        formatted = f"{d[:2]}.{d[2:5]}.{d[5:8]}/{d[8:]}"
                    else:
                        formatted = f"{d[:2]}.{d[2:5]}.{d[5:8]}/{d[8:12]}-{d[12:]}"
            elif kind == "date":
                # Data: dd/mm/aaaa
                digits_trim = digits[:8]
                if len(digits_trim) <= 2:
                    formatted = digits_trim
                elif len(digits_trim) <= 4:
                    formatted = f"{digits_trim[:2]}/{digits_trim[2:4]}"
                else:
                    formatted = f"{digits_trim[:2]}/{digits_trim[2:4]}/{digits_trim[4:8]}"
            elif kind == "time":
                # Hora: HH:MM
                digits_trim = digits[:4]
                if len(digits_trim) <= 2:
                    formatted = digits_trim
                else:
                    formatted = f"{digits_trim[:2]}:{digits_trim[2:]}"
            elif kind == "money":
                # Dinheiro: R$ #.###,##
                if not digits:
                    formatted = ""
                else:
                    # limite de segurança
                    digits_trim = digits[:15]

                    # separa centavos (2 últimos dígitos)
                    if len(digits_trim) == 1:
                        inteiro = "0"
                        centavos = f"0{digits_trim}"
                    elif len(digits_trim) == 2:
                        inteiro = "0"
                        centavos = digits_trim
                    else:
                        inteiro = digits_trim[:-2]
                        centavos = digits_trim[-2:]

                    inteiro = inteiro.lstrip("0") or "0"

                    grupos = []
                    while len(inteiro) > 3:
                        grupos.insert(0, inteiro[-3:])
                        inteiro = inteiro[:-3]
                    grupos.insert(0, inteiro)
                    inteiro_fmt = ".".join(grupos)

                    formatted = f"R$ {inteiro_fmt},{centavos}"
            else:
                formatted = text  # nenhuma máscara

            entry.delete(0, "end")
            entry.insert(0, formatted)

        entry.bind("<KeyRelease>", on_key_release)

    # ----------------- CONTRATANTE -----------------
    def _build_tab_contratante(self, parent: ctk.CTkFrame):
        frame = ctk.CTkScrollableFrame(parent)
        frame.pack(fill="both", expand=True, padx=10, pady=10)

        # Tipo
        ctk.CTkLabel(frame, text="Tipo de contratante").grid(row=0, column=0, sticky="w", pady=(0, 5))
        self.inputs["contratante_tipo"] = ctk.CTkComboBox(
            frame,
            values=["Pessoa Física", "Pessoa Jurídica"],
            width=180
        )
        self.inputs["contratante_tipo"].set("Pessoa Física")
        self.inputs["contratante_tipo"].grid(row=1, column=0, sticky="w")

        # Linha
        ctk.CTkLabel(frame, text="").grid(row=2, column=0, pady=5)

        # Dados básicos
        self._add_labeled_entry(frame, "Nome / Razão Social", "contratante_nome_razao", row=3)
        self._add_labeled_entry(frame, "CPF / CNPJ", "contratante_cpf_cnpj", row=4, width=180)
        self._add_labeled_entry(frame, "Telefone", "contratante_telefone", row=5, width=180)
        self._add_labeled_entry(frame, "E-mail", "contratante_email", row=6, width=300)

        # Endereço
        ctk.CTkLabel(frame, text="Endereço", font=ctk.CTkFont(weight="bold")).grid(
            row=7, column=0, sticky="w", pady=(15, 5)
        )
        self._add_labeled_entry(frame, "Logradouro", "contratante_endereco_logradouro", row=8, width=350)

        # Número + complemento
        row = 9
        ctk.CTkLabel(frame, text="Número", width=130, anchor="w").grid(row=row, column=0, sticky="w")
        entry_num = ctk.CTkEntry(frame, width=80)
        entry_num.grid(row=row, column=1, sticky="w", padx=(0, 10))
        self.inputs["contratante_endereco_numero"] = entry_num

        ctk.CTkLabel(frame, text="Compl.", width=60, anchor="w").grid(row=row, column=2, sticky="w")
        entry_comp = ctk.CTkEntry(frame, width=100)
        entry_comp.grid(row=row, column=3, sticky="w")
        self.inputs["contratante_endereco_complemento"] = entry_comp

        self._add_labeled_entry(frame, "Bairro", "contratante_endereco_bairro", row=10, width=250)

        # Cidade / UF
        row = 11
        ctk.CTkLabel(frame, text="Cidade", width=130, anchor="w").grid(row=row, column=0, sticky="w")
        entry_cid = ctk.CTkEntry(frame, width=250)
        entry_cid.grid(row=row, column=1, sticky="w", padx=(0, 10))
        self.inputs["contratante_endereco_cidade"] = entry_cid

        ctk.CTkLabel(frame, text="UF", width=30, anchor="w").grid(row=row, column=2, sticky="w")
        entry_uf = ctk.CTkEntry(frame, width=40)
        entry_uf.grid(row=row, column=3, sticky="w")
        self.inputs["contratante_endereco_uf"] = entry_uf

        # CEP + botão Buscar
        row = 12
        ctk.CTkLabel(frame, text="CEP", width=130, anchor="w").grid(row=row, column=0, sticky="w", pady=3)
        entry_cep_contr = ctk.CTkEntry(frame, width=120)
        entry_cep_contr.grid(row=row, column=1, sticky="w", pady=3)
        self.inputs["contratante_endereco_cep"] = entry_cep_contr

        btn_buscar_cep_contr = ctk.CTkButton(
            frame,
            text="Buscar",
            width=80,
            command=self.buscar_cep_contratante
        )
        btn_buscar_cep_contr.grid(row=row, column=2, sticky="w", padx=(10, 0), pady=3)

        # Representante
        ctk.CTkLabel(frame, text="Dados do representante (apenas PJ)",
                     font=ctk.CTkFont(weight="bold")).grid(
            row=13, column=0, sticky="w", pady=(15, 5)
        )
        self._add_labeled_entry(frame, "Nome", "contratante_representante_nome", row=14, width=350)
        self._add_labeled_entry(frame, "CPF", "contratante_representante_cpf", row=15, width=180)

        # Config grid
        for col in range(4):
            frame.grid_columnconfigure(col, weight=0)
        frame.grid_columnconfigure(1, weight=1)

    # ----------------- CONTRATADO -----------------
    def _build_tab_contratado(self, parent: ctk.CTkFrame):
        frame = ctk.CTkScrollableFrame(parent)
        frame.pack(fill="both", expand=True, padx=10, pady=10)

        # Tipo
        ctk.CTkLabel(frame, text="Tipo de contratado").grid(row=0, column=0, sticky="w", pady=(0, 5))
        self.inputs["contratado_tipo"] = ctk.CTkComboBox(
            frame,
            values=["Pessoa Física", "Pessoa Jurídica"],
            width=180
        )
        self.inputs["contratado_tipo"].set("Pessoa Jurídica")
        self.inputs["contratado_tipo"].grid(row=1, column=0, sticky="w")

        # Linha
        ctk.CTkLabel(frame, text="").grid(row=2, column=0, pady=5)

        # Dados básicos
        self._add_labeled_entry(frame, "Nome / Razão Social", "contratado_nome_razao", row=3)
        self._add_labeled_entry(frame, "CPF / CNPJ", "contratado_cpf_cnpj", row=4, width=180)
        self._add_labeled_entry(frame, "Telefone", "contratado_telefone", row=5, width=180)
        self._add_labeled_entry(frame, "E-mail", "contratado_email", row=6, width=300)

        # Endereço
        ctk.CTkLabel(frame, text="Endereço", font=ctk.CTkFont(weight="bold")).grid(
            row=7, column=0, sticky="w", pady=(15, 5)
        )
        self._add_labeled_entry(frame, "Logradouro", "contratado_endereco_logradouro", row=8, width=350)

        # Número + complemento
        row = 9
        ctk.CTkLabel(frame, text="Número", width=130, anchor="w").grid(row=row, column=0, sticky="w")
        entry_num = ctk.CTkEntry(frame, width=80)
        entry_num.grid(row=row, column=1, sticky="w", padx=(0, 10))
        self.inputs["contratado_endereco_numero"] = entry_num

        ctk.CTkLabel(frame, text="Compl.", width=60, anchor="w").grid(row=row, column=2, sticky="w")
        entry_comp = ctk.CTkEntry(frame, width=100)
        entry_comp.grid(row=row, column=3, sticky="w")
        self.inputs["contratado_endereco_complemento"] = entry_comp

        self._add_labeled_entry(frame, "Bairro", "contratado_endereco_bairro", row=10, width=250)

        # Cidade / UF
        row = 11
        ctk.CTkLabel(frame, text="Cidade", width=130, anchor="w").grid(row=row, column=0, sticky="w")
        entry_cid = ctk.CTkEntry(frame, width=250)
        entry_cid.grid(row=row, column=1, sticky="w", padx=(0, 10))
        self.inputs["contratado_endereco_cidade"] = entry_cid

        ctk.CTkLabel(frame, text="UF", width=30, anchor="w").grid(row=row, column=2, sticky="w")
        entry_uf = ctk.CTkEntry(frame, width=40)
        entry_uf.grid(row=row, column=3, sticky="w")
        self.inputs["contratado_endereco_uf"] = entry_uf

        # CEP + botão Buscar (contratado)
        row = 12
        ctk.CTkLabel(frame, text="CEP", width=130, anchor="w").grid(row=row, column=0, sticky="w", pady=3)
        entry_cep_contratado = ctk.CTkEntry(frame, width=120)
        entry_cep_contratado.grid(row=row, column=1, sticky="w", pady=3)
        self.inputs["contratado_endereco_cep"] = entry_cep_contratado

        btn_buscar_cep_contratado = ctk.CTkButton(
            frame,
            text="Buscar",
            width=80,
            command=self.buscar_cep_contratado
        )
        btn_buscar_cep_contratado.grid(row=row, column=2, sticky="w", padx=(10, 0), pady=3)

        # Representante
        ctk.CTkLabel(frame, text="Representante (apenas PJ)",
                     font=ctk.CTkFont(weight="bold")).grid(
            row=13, column=0, sticky="w", pady=(15, 5)
        )
        self._add_labeled_entry(frame, "Nome", "contratado_representante_nome", row=14, width=350)
        self._add_labeled_entry(frame, "CPF", "contratado_representante_cpf", row=15, width=180)

        for col in range(4):
            frame.grid_columnconfigure(col, weight=0)
        frame.grid_columnconfigure(1, weight=1)

    # ----------------- EVENTO / LOCAL -----------------
    def _build_tab_evento(self, parent: ctk.CTkFrame):
        frame = ctk.CTkScrollableFrame(parent)
        frame.pack(fill="both", expand=True, padx=10, pady=10)

        ctk.CTkLabel(frame, text="Dados do evento", font=ctk.CTkFont(weight="bold")).grid(
            row=0, column=0, sticky="w", pady=(0, 5)
        )

        self._add_labeled_entry(frame, "Nome do evento", "evento_nome", row=1, width=350)
        self._add_labeled_entry(frame, "Atração musical", "evento_atracao_musical", row=2, width=350)
        self._add_labeled_entry(frame, "Data (dd/mm/aaaa)", "evento_data", row=3, width=120)
        self._add_labeled_entry(frame, "Horário início (hh:mm)", "evento_horario_inicio", row=4, width=80)
        self._add_labeled_entry(frame, "Horário fim (prev.)", "evento_horario_fim_previsto", row=5, width=80)

        ctk.CTkLabel(frame, text="Local da apresentação", font=ctk.CTkFont(weight="bold")).grid(
            row=6, column=0, sticky="w", pady=(15, 5)
        )

        self._add_labeled_entry(frame, "Nome do local", "evento_local_nome", row=7, width=350)
        self._add_labeled_entry(frame, "Logradouro", "evento_local_logradouro", row=8, width=350)

        # Número + compl.
        row = 9
        ctk.CTkLabel(frame, text="Número", width=130, anchor="w").grid(row=row, column=0, sticky="w")
        entry_num = ctk.CTkEntry(frame, width=80)
        entry_num.grid(row=row, column=1, sticky="w", padx=(0, 10))
        self.inputs["evento_local_numero"] = entry_num

        ctk.CTkLabel(frame, text="Compl.", width=60, anchor="w").grid(row=row, column=2, sticky="w")
        entry_comp = ctk.CTkEntry(frame, width=100)
        entry_comp.grid(row=row, column=3, sticky="w")
        self.inputs["evento_local_complemento"] = entry_comp

        self._add_labeled_entry(frame, "Bairro", "evento_local_bairro", row=10, width=250)

        # Cidade / UF
        row = 11
        ctk.CTkLabel(frame, text="Cidade", width=130, anchor="w").grid(row=row, column=0, sticky="w")
        entry_cid = ctk.CTkEntry(frame, width=250)
        entry_cid.grid(row=row, column=1, sticky="w", padx=(0, 10))
        self.inputs["evento_local_cidade"] = entry_cid

        ctk.CTkLabel(frame, text="UF", width=30, anchor="w").grid(row=row, column=2, sticky="w")
        entry_uf = ctk.CTkEntry(frame, width=40)
        entry_uf.grid(row=row, column=3, sticky="w")
        self.inputs["evento_local_uf"] = entry_uf

        # CEP + buscar
        row = 12
        ctk.CTkLabel(frame, text="CEP", width=130, anchor="w").grid(row=row, column=0, sticky="w", pady=3)
        entry_cep_evento = ctk.CTkEntry(frame, width=120)
        entry_cep_evento.grid(row=row, column=1, sticky="w", pady=3)
        self.inputs["evento_local_cep"] = entry_cep_evento

        btn_buscar_cep_evento = ctk.CTkButton(
            frame,
            text="Buscar",
            width=80,
            command=self.buscar_cep_evento
        )
        btn_buscar_cep_evento.grid(row=row, column=2, sticky="w", padx=(10, 0), pady=3)

        for col in range(4):
            frame.grid_columnconfigure(col, weight=0)
        frame.grid_columnconfigure(1, weight=1)

    # ----------------- SOM / ALIMENTAÇÃO -----------------
    def _build_tab_som(self, parent: ctk.CTkFrame):
        frame = ctk.CTkFrame(parent)
        frame.pack(fill="both", expand=True, padx=10, pady=10)

        ctk.CTkLabel(frame, text="Responsabilidade pelo som",
                     font=ctk.CTkFont(weight="bold")).pack(anchor="w", pady=(0, 10))

        ctk.CTkLabel(frame, text="Quem fornece o som (PA, caixas, etc.)?").pack(anchor="w", pady=(0, 5))

        rb_frame = ctk.CTkFrame(frame)
        rb_frame.pack(anchor="w", pady=5)

        ctk.CTkRadioButton(
            rb_frame,
            text="Banda",
            variable=self.som_responsavel_var,
            value="Banda"
        ).pack(side="left", padx=(0, 15))

        ctk.CTkRadioButton(
            rb_frame,
            text="Contratante",
            variable=self.som_responsavel_var,
            value="Contratante"
        ).pack(side="left")

        ctk.CTkLabel(
            frame,
            text="Obs.: o texto da cláusula será gerado automaticamente com base nessa escolha.",
            wraplength=600,
            justify="left"
        ).pack(anchor="w", pady=(10, 0))
        
            # Alimentação (cláusula opcional)
        ctk.CTkLabel(
            frame,
            text="Alimentação",
            font=ctk.CTkFont(weight="bold")
        ).pack(anchor="w", pady=(20, 5))

        ctk.CTkLabel(
            frame,
            text="Deseja incluir cláusula de alimentação/consumação para o staff da banda?"
        ).pack(anchor="w")

        alimentacao_frame = ctk.CTkFrame(frame)
        alimentacao_frame.pack(anchor="w", pady=5)

        ctk.CTkRadioButton(
            alimentacao_frame,
            text="Sim",
            variable=self.alimentacao_var,
            value="Sim"
        ).pack(side="left", padx=(0, 15))

        ctk.CTkRadioButton(
            alimentacao_frame,
            text="Não",
            variable=self.alimentacao_var,
            value="Não"
        ).pack(side="left")

    # ----------------- PAGAMENTO -----------------
    def _build_tab_pagamento(self, parent: ctk.CTkFrame):
        frame = ctk.CTkScrollableFrame(parent)
        frame.pack(fill="both", expand=True, padx=10, pady=10)

        ctk.CTkLabel(
            frame,
            text="Valor e forma de pagamento",
            font=ctk.CTkFont(weight="bold")
        ).grid(row=0, column=0, sticky="w", pady=(0, 5))

        # Valor total
        self._add_labeled_entry(frame, "Valor total (R$)", "pagamento_valor_total", row=1, width=150)

        # Forma
        ctk.CTkLabel(frame, text="Forma", width=130, anchor="w").grid(row=2, column=0, sticky="w", pady=(5, 5))
        cb_forma = ctk.CTkComboBox(
            frame,
            values=["À vista", "Sinal + restante", "Parcelado", "Outro"],
            width=180,
            command=self._on_pagamento_forma_change
        )
        cb_forma.set("À vista")
        cb_forma.grid(row=2, column=1, sticky="w")
        self.inputs["pagamento_forma"] = cb_forma

        # Meio
        ctk.CTkLabel(frame, text="Meio", width=130, anchor="w").grid(row=3, column=0, sticky="w", pady=(5, 5))
        cb_meio = ctk.CTkComboBox(
            frame,
            values=["PIX", "TED/DOC", "Dinheiro", "Boleto", "Cartão", "Outro"],
            width=180
        )
        cb_meio.set("PIX")
        cb_meio.grid(row=3, column=1, sticky="w")
        self.inputs["pagamento_meio"] = cb_meio

        # --- Seções específicas por forma de pagamento ---

        # À vista
        self.pag_frame_avista = ctk.CTkFrame(frame)
        self.pag_frame_avista.grid(row=4, column=0, columnspan=4, sticky="w", pady=(15, 5))
        ctk.CTkLabel(
            self.pag_frame_avista,
            text="À vista",
            font=ctk.CTkFont(weight="bold")
        ).grid(row=0, column=0, sticky="w", pady=(0, 5))
        self._add_labeled_entry(
            self.pag_frame_avista,
            "Data de pagamento (dd/mm/aaaa)",
            "pagamento_data_unica",
            row=1,
            width=150
        )

        # Sinal + restante
        self.pag_frame_sinal = ctk.CTkFrame(frame)
        self.pag_frame_sinal.grid(row=5, column=0, columnspan=4, sticky="w", pady=(15, 5))
        ctk.CTkLabel(
            self.pag_frame_sinal,
            text="Sinal + restante",
            font=ctk.CTkFont(weight="bold")
        ).grid(row=0, column=0, sticky="w", pady=(0, 5))
        self._add_labeled_entry(
            self.pag_frame_sinal,
            "Sinal (%)",
            "pagamento_sinal_percentual",
            row=1,
            width=80
        )
        self._add_labeled_entry(
            self.pag_frame_sinal,
            "Data do sinal (dd/mm/aaaa)",
            "pagamento_sinal_data",
            row=2,
            width=150
        )
        self._add_labeled_entry(
            self.pag_frame_sinal,
            "Data do restante (dd/mm/aaaa)",
            "pagamento_restante_data",
            row=3,
            width=150
        )

        # Parcelado
        self.pag_frame_parc = ctk.CTkFrame(frame)
        self.pag_frame_parc.grid(row=6, column=0, columnspan=4, sticky="w", pady=(15, 5))
        ctk.CTkLabel(
            self.pag_frame_parc,
            text="Parcelado (MVP)",
            font=ctk.CTkFont(weight="bold")
        ).grid(row=0, column=0, sticky="w", pady=(0, 5))
        self._add_labeled_entry(
            self.pag_frame_parc,
            "Nº de parcelas",
            "pagamento_num_parcelas",
            row=1,
            width=80
        )
        self._add_labeled_entry(
            self.pag_frame_parc,
            "1ª parcela em (dd/mm/aaaa)",
            "pagamento_primeira_parcela_data",
            row=2,
            width=150
        )

        ctk.CTkLabel(self.pag_frame_parc, text="Periodicidade", width=130, anchor="w").grid(
            row=3, column=0, sticky="w", pady=(5, 5)
        )
        cb_period = ctk.CTkComboBox(
            self.pag_frame_parc,
            values=["Mensal", "Semanal", "Outro"],
            width=120
        )
        cb_period.set("Mensal")
        cb_period.grid(row=3, column=1, sticky="w")
        self.inputs["pagamento_periodicidade"] = cb_period

        for col in range(4):
            frame.grid_columnconfigure(col, weight=0)
        frame.grid_columnconfigure(1, weight=1)

    def _on_pagamento_forma_change(self, choice: str):
        """Callback chamado ao mudar a forma de pagamento no ComboBox."""
        self._update_pagamento_forma_ui(choice)

    def _update_pagamento_forma_ui(self, forma: str | None = None):
        """Mostra apenas a seção de campos correspondente à forma de pagamento escolhida."""
        if forma is None:
            cb = self.inputs.get("pagamento_forma")
            try:
                forma = cb.get() if cb else ""
            except Exception:
                forma = ""
        forma = (forma or "").strip()

        # Garante que os frames existem antes de manipular
        frames = [
            f for f in (self.pag_frame_avista, self.pag_frame_sinal, self.pag_frame_parc)
            if f is not None
        ]
        if not frames:
            return

        # Esconde todas as seções
        for f in frames:
            f.grid_remove()

        # Mostra apenas a correspondente
        if forma == "À vista":
            if self.pag_frame_avista is not None:
                self.pag_frame_avista.grid()
        elif forma == "Sinal + restante":
            if self.pag_frame_sinal is not None:
                self.pag_frame_sinal.grid()
        elif forma == "Parcelado":
            if self.pag_frame_parc is not None:
                self.pag_frame_parc.grid()
        else:
            # "Outro" -> nenhuma seção específica
            pass

    # ----------------- FAVORECIDO -----------------
    def _build_tab_favorecido(self, parent: ctk.CTkFrame):
        frame = ctk.CTkScrollableFrame(parent)
        frame.pack(fill="both", expand=True, padx=10, pady=10)

        ctk.CTkLabel(frame, text="Dados do favorecido",
                font=ctk.CTkFont(weight="bold")).grid(row=0, column=0, sticky="w", pady=(0, 5))

        self._add_labeled_entry(frame, "Nome", "favorecido_nome", row=1, width=350)
        self._add_labeled_entry(frame, "CPF / CNPJ", "favorecido_cpf_cnpj", row=2, width=180)

        # checkbox para copiar dados do CONTRATADO
        ctk.CTkCheckBox(
            frame,
            text="Favorecido é o mesmo que o CONTRATADO",
            variable=self.favorecido_igual_contratado_var,
            command=self._on_toggle_favorecido_igual_contratado
        ).grid(row=3, column=0, sticky="w", pady=(5, 10))

        ctk.CTkLabel(frame, text="Dados bancários (opcional)",
                     font=ctk.CTkFont(weight="bold")).grid(row=4, column=0, sticky="w", pady=(15, 5))

        self._add_labeled_entry(frame, "Banco", "favorecido_banco_nome", row=5, width=250)
        self._add_labeled_entry(frame, "Cód. banco", "favorecido_banco_codigo", row=6, width=80)
        self._add_labeled_entry(frame, "Agência", "favorecido_agencia", row=7, width=100)
        self._add_labeled_entry(frame, "Conta", "favorecido_conta", row=8, width=140)

        ctk.CTkLabel(frame, text="Tipo conta", width=130, anchor="w").grid(row=9, column=0, sticky="w", pady=(5, 5))
        cb_tipoconta = ctk.CTkComboBox(
            frame,
            values=["Corrente", "Poupança", "Pagamento"],
            width=150
        )
        cb_tipoconta.grid(row=9, column=1, sticky="w")
        self.inputs["favorecido_tipo_conta"] = cb_tipoconta

        ctk.CTkLabel(frame, text="PIX", font=ctk.CTkFont(weight="bold")).grid(
            row=10, column=0, sticky="w", pady=(15, 5)
        )

        self._add_labeled_entry(frame, "Chave PIX", "favorecido_pix_chave", row=11, width=350)

        ctk.CTkLabel(frame, text="Tipo da chave", width=130, anchor="w").grid(row=12, column=0, sticky="w", pady=(5, 5))
        cb_pixtipo = ctk.CTkComboBox(
            frame,
            values=["CPF/CNPJ", "E-mail", "Telefone", "Chave aleatória"],
            width=150
        )
        cb_pixtipo.grid(row=12, column=1, sticky="w")
        self.inputs["favorecido_pix_tipo"] = cb_pixtipo

        for col in range(4):
            frame.grid_columnconfigure(col, weight=0)
        frame.grid_columnconfigure(1, weight=1)

    def _on_toggle_favorecido_igual_contratado(self):
        """Quando marcado, copia nome/CPF do CONTRATADO para o FAVORECIDO e bloqueia edição."""
        marcado = self.favorecido_igual_contratado_var.get()

        fav_nome = self.inputs.get("favorecido_nome")
        fav_doc = self.inputs.get("favorecido_cpf_cnpj")
        if not fav_nome or not fav_doc:
            return

        if marcado:
            contratado_nome_widget = self.inputs.get("contratado_nome_razao")
            contratado_doc_widget = self.inputs.get("contratado_cpf_cnpj")

            nome = contratado_nome_widget.get() if contratado_nome_widget else ""
            doc = contratado_doc_widget.get() if contratado_doc_widget else ""

            fav_nome.configure(state="normal")
            fav_doc.configure(state="normal")

            fav_nome.delete(0, "end")
            fav_nome.insert(0, nome)

            fav_doc.delete(0, "end")
            fav_doc.insert(0, doc)

            fav_nome.configure(state="disabled")
            fav_doc.configure(state="disabled")
        else:
            fav_nome.configure(state="normal")
            fav_doc.configure(state="normal")
            fav_nome.delete(0, "end")
            fav_doc.delete(0, "end")
            
    # ----------------- RESUMO -----------------
    def _build_tab_resumo(self, parent: ctk.CTkFrame):
        frame = ctk.CTkFrame(parent)
        frame.pack(fill="both", expand=True, padx=10, pady=10)

        ctk.CTkLabel(frame, text="Resumo do contrato",
                     font=ctk.CTkFont(weight="bold")).pack(anchor="w", pady=(0, 5))

        btn_atualizar_resumo = ctk.CTkButton(
            frame,
            text="Atualizar resumo",
            command=self._update_resumo_preview
        )
        btn_atualizar_resumo.pack(anchor="w", pady=(0, 5))

        self.preview_box = ctk.CTkTextbox(frame, width=900, height=450)
        self.preview_box.pack(fill="both", expand=True, pady=(5, 0))

    # ---------------------------------------------------------
    # Utilitários de layout
    # ---------------------------------------------------------
    def _add_labeled_entry(self, parent, label, key, row, width=300):
        ctk.CTkLabel(parent, text=label, width=130, anchor="w").grid(row=row, column=0, sticky="w", pady=3)
        entry = ctk.CTkEntry(parent, width=width)
        entry.grid(row=row, column=1, sticky="w", pady=3)
        self.inputs[key] = entry

    # ---------------------------------------------------------
    # Busca CEP (ViaCEP)
    # ---------------------------------------------------------
    def _preencher_endereco_por_cep(self, cep_key: str, logradouro_key: str,
                                    bairro_key: str, cidade_key: str, uf_key: str):
        """Usa o CEP (em self.inputs[cep_key]) para preencher logradouro/bairro/cidade/UF."""
        cep_widget = self.inputs.get(cep_key)
        if not cep_widget:
            return

        cep_raw = cep_widget.get() or ""
        cep = "".join(c for c in cep_raw if c.isdigit())

        if len(cep) != 8:
            messagebox.showerror("CEP inválido", "Informe um CEP com 8 dígitos.")
            return

        try:
            resp = requests.get(f"https://viacep.com.br/ws/{cep}/json/", timeout=5)
            resp.raise_for_status()
            data = resp.json()
        except Exception as e:
            messagebox.showerror("Erro na consulta",
                                 f"Não foi possível consultar o CEP.\n\nDetalhes: {e}")
            return

        if data.get("erro"):
            messagebox.showerror("CEP não encontrado",
                                 "Não foi possível localizar o CEP informado.")
            return

        # Preenche campos
        def set_input(key, value):
            widget = self.inputs.get(key)
            if widget and value:
                try:
                    widget.delete(0, "end")
                    widget.insert(0, value)
                except Exception:
                    pass

        set_input(logradouro_key, data.get("logradouro", ""))
        set_input(bairro_key, data.get("bairro", ""))
        set_input(cidade_key, data.get("localidade", ""))
        set_input(uf_key, data.get("uf", ""))

    def buscar_cep_contratante(self):
        self._preencher_endereco_por_cep(
            cep_key="contratante_endereco_cep",
            logradouro_key="contratante_endereco_logradouro",
            bairro_key="contratante_endereco_bairro",
            cidade_key="contratante_endereco_cidade",
            uf_key="contratante_endereco_uf",
        )

    def buscar_cep_evento(self):
        self._preencher_endereco_por_cep(
            cep_key="evento_local_cep",
            logradouro_key="evento_local_logradouro",
            bairro_key="evento_local_bairro",
            cidade_key="evento_local_cidade",
            uf_key="evento_local_uf",
        )

    def buscar_cep_contratado(self):
        self._preencher_endereco_por_cep(
            cep_key="contratado_endereco_cep",
            logradouro_key="contratado_endereco_logradouro",
            bairro_key="contratado_endereco_bairro",
            cidade_key="contratado_endereco_cidade",
            uf_key="contratado_endereco_uf",
        )

    # ---------------------------------------------------------
    # Lógica de botões
    # ---------------------------------------------------------
    def limpar_campos(self):
        for widget in self.inputs.values():
            # Entry e ComboBox têm .delete; ComboBox também aceita set("")
            try:
                widget.delete(0, "end")
            except Exception:
                pass
            try:
                widget.set("")
            except Exception:
                pass

        self.som_responsavel_var.set("Contratante")
        self.alimentacao_var.set("Sim")
        self.preview_box.delete("1.0", "end")
        
        self.favorecido_igual_contratado_var.set(False)
        self._on_toggle_favorecido_igual_contratado()
        try:
            self.inputs["pagamento_forma"].set("À vista")
        except Exception:
            pass
        self._update_pagamento_forma_ui("À vista")

    def _update_resumo_preview(self):
        """Atualiza o resumo sem gerar o contrato."""
        values = {k: (widget.get() if hasattr(widget, "get") else "") for k, widget in self.inputs.items()}
        som = self.som_responsavel_var.get()
        alimentacao = self.alimentacao_var.get()

        resumo = []

        resumo.append("CONTRATO DE PRESTAÇÃO DE SERVIÇOS MUSICAIS\n")
        resumo.append("-" * 60 + "\n\n")

        resumo.append("CONTRATANTE:\n")
        resumo.append(f"  Nome/Razão Social: {values.get('contratante_nome_razao', '')}\n")
        resumo.append(f"  CPF/CNPJ: {values.get('contratante_cpf_cnpj', '')}\n")
        resumo.append(f"  Endereço: {values.get('contratante_endereco_logradouro', '')}, "
                      f"{values.get('contratante_endereco_numero', '')} - "
                      f"{values.get('contratante_endereco_bairro', '')}, "
                      f"{values.get('contratante_endereco_cidade', '')}/"
                      f"{values.get('contratante_endereco_uf', '')} - "
                      f"CEP: {values.get('contratante_endereco_cep', '')}\n")
        resumo.append(f"  Telefone: {values.get('contratante_telefone', '')}\n")
        resumo.append(f"  E-mail: {values.get('contratante_email', '')}\n\n")

        resumo.append("CONTRATADO:\n")
        resumo.append(f"  Nome/Razão Social: {values.get('contratado_nome_razao', '')}\n")
        resumo.append(f"  CPF/CNPJ: {values.get('contratado_cpf_cnpj', '')}\n")
        resumo.append(f"  Endereço: {values.get('contratado_endereco_logradouro', '')}, "
                      f"{values.get('contratado_endereco_numero', '')} - "
                      f"{values.get('contratado_endereco_bairro', '')}, "
                      f"{values.get('contratado_endereco_cidade', '')}/"
                      f"{values.get('contratado_endereco_uf', '')} - "
                      f"CEP: {values.get('contratado_endereco_cep', '')}\n")
        resumo.append(f"  Telefone: {values.get('contratado_telefone', '')}\n")
        resumo.append(f"  E-mail: {values.get('contratado_email', '')}\n\n")

        resumo.append("EVENTO:\n")
        resumo.append(f"  Nome do evento: {values.get('evento_nome', '')}\n")
        resumo.append(f"  Data: {values.get('evento_data', '')}\n")
        resumo.append(
            f"  Horário: {values.get('evento_horario_inicio', '')}h às "
            f"{values.get('evento_horario_fim_previsto', '')}h\n"
        )
        resumo.append(
            "  Local: "
            f"{values.get('evento_local_nome', '')}, "
            f"{values.get('evento_local_logradouro', '')}, "
            f"{values.get('evento_local_numero', '')} - "
            f"{values.get('evento_local_bairro', '')}, "
            f"{values.get('evento_local_cidade', '')}/"
            f"{values.get('evento_local_uf', '')} - "
            f"CEP: {values.get('evento_local_cep', '')}\n\n"
        )

        resumo.append("RESPONSABILIDADE PELO SOM:\n")
        if som == "Banda":
            resumo.append("  A banda será responsável por levar e operar o sistema de som necessário.\n\n")
        else:
            resumo.append("  O CONTRATANTE será responsável pelo sistema de som necessário.\n\n")

        resumo.append("ALIMENTAÇÃO:\n")
        if alimentacao == "Sim":
            resumo.append("  Haverá fornecimento de alimentação/consumação ao staff.\n\n")
        else:
            resumo.append("  Não haverá fornecimento de alimentação.\n\n")

        resumo.append("PAGAMENTO:\n")
        resumo.append(f"  Valor total: R$ {values.get('pagamento_valor_total', '')}\n")
        resumo.append(f"  Forma: {values.get('pagamento_forma', '')}\n")
        resumo.append(f"  Meio: {values.get('pagamento_meio', '')}\n\n")

        resumo.append("FAVORECIDO:\n")
        resumo.append(f"  Nome: {values.get('favorecido_nome', '')}\n")
        resumo.append(f"  CPF/CNPJ: {values.get('favorecido_cpf_cnpj', '')}\n")
        resumo.append(
            f"  Chave PIX: {values.get('favorecido_pix_chave', '')} "
            f"({values.get('favorecido_pix_tipo', '')})\n\n"
        )

        resumo.append("(Resumo prévio — o contrato completo será gerado ao clicar em 'Gerar contrato'.)\n")

        self.preview_box.delete("1.0", "end")
        self.preview_box.insert("1.0", "".join(resumo))

    def _on_tab_change(self):
        """Callback do TabView — detecta a aba ativa e atualiza o resumo se for a aba Resumo."""
        try:
            aba = self.tabview.get()
            if aba == "Resumo":
                self._update_resumo_preview()
        except Exception:
            pass

    def gerar_contrato(self):
        """MVP: monta um texto de resumo com base em alguns campos, mostra na aba Resumo e gera um DOCX."""
        # coleta de dados
        values = {}
        for key, widget in self.inputs.items():
            try:
                values[key] = widget.get()
            except Exception:
                values[key] = ""

        som = self.som_responsavel_var.get()
        alimentacao = self.alimentacao_var.get()
        
        snapshot = {
            "values": values,
            "som": som,
            "alimentacao": alimentacao,
            "favorecido_igual_contratado": self.favorecido_igual_contratado_var.get(),
        }

        # ---------- PREVIEW ----------
        self._update_resumo_preview()

        # ---------- GERAÇÃO DO DOCX ----------
        try:
            contexto = montar_contexto(values, som, alimentacao)

            # usa template único; cláusula de som é preenchida via placeholder
            template = TEMPLATE_CONTRATO

            if not template.exists():
                raise FileNotFoundError(f"Template não encontrado: {template}")

            # nome base sem extensão (data: dd/mm/yyyy -> yyyymmdd)
            evento_raw = values.get("evento_data", "")
            m = re.match(r"^\s*(\d{2})/(\d{2})/(\d{4})\s*$", evento_raw)
            if m:
                data_evento = f"{m.group(3)}{m.group(2)}{m.group(1)}"
            else:
                data_evento = evento_raw.replace("/", "")
            atracao = values.get("evento_atracao_musical", "").replace(" ", "_")
            base_name = f"Contrato_{atracao}_{data_evento}"

            # --- Versionamento automático v1, v2, v3... ---
            versao = 1
            while True:
                arquivo_saida = SAIDA_DIR / f"{base_name}_v{versao}.docx"
                json_path = SAIDA_DIR / f"{base_name}_v{versao}.json"
                if not arquivo_saida.exists() and not json_path.exists():
                    break
                versao += 1

            # adiciona a versão também dentro do JSON
            snapshot["versao"] = versao

            import json
            with open(json_path, "w", encoding="utf-8") as f:
                json.dump(snapshot, f, ensure_ascii=False, indent=2)

            preencher_template_docx(template, arquivo_saida, contexto)

            messagebox.showinfo(
                "Contrato gerado",
                f"Contrato gerado com sucesso em:\n{arquivo_saida}"
            )
        except Exception as e:
            messagebox.showerror(
                "Erro ao gerar contrato",
                f"Ocorreu um erro ao gerar o contrato:\n{e}"
            )

    def carregar_preenchimento(self):
        path = filedialog.askopenfilename(
            title="Selecione o preenchimento do contrato",
            initialdir=SAIDA_DIR,
            filetypes=[("JSON", "*.json")]
        )
        if not path:
            return

        import json
        with open(path, "r", encoding="utf-8") as f:
            snapshot = json.load(f)

        values = snapshot.get("values", {})
        som = snapshot.get("som", "Contratante")
        alimentacao = snapshot.get("alimentacao", "Não")
        fav_igual = snapshot.get("favorecido_igual_contratado", False)

        # repopula os campos
        for key, value in values.items():
            widget = self.inputs.get(key)
            if not widget:
                continue
            try:
                widget.delete(0, "end")
                widget.insert(0, value)
            except Exception:
                # ComboBox etc.
                try:
                    widget.set(value)
                except Exception:
                    pass

        # repopula radios
        self.som_responsavel_var.set(som)
        self.alimentacao_var.set(alimentacao)
        self.favorecido_igual_contratado_var.set(bool(fav_igual))
        self._on_toggle_favorecido_igual_contratado()
        # Atualiza visibilidade das seções de pagamento com base na forma carregada
        try:
            forma_loaded = values.get("pagamento_forma", "")
        except Exception:
            forma_loaded = ""
        self._update_pagamento_forma_ui(forma_loaded)
        # Atualiza o resumo se já estiver na aba Resumo
        try:
            if self.tabview.get() == "Resumo":
                self._update_resumo_preview()
        except Exception:
            pass

if __name__ == "__main__":
    print("Iniciando ContractApp...")
    app = ContractApp()
    print("Entrando no mainloop...")
    app.mainloop()
    print("Saiu do mainloop")